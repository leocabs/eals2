import pandas as pd
import joblib
from sklearn.svm import SVC
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import accuracy_score, confusion_matrix, classification_report, roc_curve, auc
from collections import Counter
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns


# Dataset path
df = pd.read_csv("data/eals-dataset-for-training.csv")

# Function Readiness
def determine_readiness(row):
    if row['level'] in ['JHS', 'SHS']:
        score_percentage = (row['flt-overall-score'] / 98) * 100
        return 1 if score_percentage >= 75 else 0
    return 0

# Apply readiness logic
df['Ready_Status'] = df.apply(determine_readiness, axis=1)

# Define features to use
features = [
    'pis_score', 'ls2slct', 'ls3-mpss', 
    'ls4-lcs', 'ls5-uss', 'ls6-dc',
    'ls1-totalEnglish', 'ls1-totalFilipino', 
    'flt-overall-score'
]


# Drop rows with missing values in selected features
df = df.dropna(subset=features + ['Ready_Status'])

# Input and target
X = df[features]
y = df['Ready_Status']

# Print original class distribution
print("\nOriginal class distribution:")
print(Counter(y))

if len(set(y)) < 2:
    print("Error: Training data must contain at least two classes.")
    exit()

# Split and scale
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Print the number of training and testing samples
print(f"Training data: {len(X_train)} samples")
print(f"Testing data: {len(X_test)} samples")

scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)


# Train model
model = SVC(probability=True, kernel='rbf')
model.fit(X_train_scaled, y_train)

# Save model and scaler
joblib.dump(model, "data/als_model.pkl")
joblib.dump(scaler, "data/als_scaler.pkl")

# Predictions and performance
y_pred = model.predict(X_test_scaled)
y_proba = model.predict_proba(X_test_scaled)[:, 1]  # Probability of class 1

accuracy = accuracy_score(y_test, y_pred)
conf_matrix = confusion_matrix(y_test, y_pred)
class_report = classification_report(y_test, y_pred)

# Cross-validation scores
cv_scores = cross_val_score(model, X, y, cv=5)
cv_mean = cv_scores.mean()
cv_std = cv_scores.std()

# ROC Curve
fpr, tpr, thresholds = roc_curve(y_test, y_proba)
roc_auc = auc(fpr, tpr)
plt.figure(figsize=(6, 4))
plt.plot(fpr, tpr, label=f'ROC Curve (AUC = {roc_auc:.2f})')
plt.plot([0, 1], [0, 1], linestyle='--', color='gray')
plt.xlabel("False Positive Rate")
plt.ylabel("True Positive Rate")
plt.title("ROC Curve")
plt.legend()
plt.tight_layout()
plt.savefig("C:\\Users\\User\\Desktop\\projects\\EALS - Copy\\backend\\models\\roc_curve.png")
plt.close()

# Confusion matrix heatmap
plt.figure(figsize=(6, 4))
sns.heatmap(conf_matrix, annot=True, fmt="d", cmap="Blues", xticklabels=["Not Ready", "Ready"], yticklabels=["Not Ready", "Ready"])
plt.title("Confusion Matrix")
plt.ylabel("Actual")
plt.xlabel("Predicted")
plt.tight_layout()
plt.savefig("C:\\Users\\User\\Desktop\\projects\\EALS - Copy\\backend\\models\\confusion_matrix.png")
plt.close()

# Ready vs Not Ready bar graph
ready_count = sum(y)
not_ready_count = len(y) - ready_count
plt.figure(figsize=(6, 4))
sns.barplot(x=["Ready", "Not Ready"], y=[ready_count, not_ready_count], hue=["Ready", "Not Ready"], dodge=False, palette="viridis", legend=False)
plt.title("Ready vs Not Ready Learners")
plt.ylabel("Number of Learners")
plt.tight_layout()
plt.savefig("C:\\Users\\User\\Desktop\\projects\\EALS - Copy\\backend\\models\\ready_not_ready_graph.png")
plt.close()

# Generate Word report
doc = Document()
doc.add_heading("ALS SVM Model Report", 0)

doc.add_heading("Feature Set", level=1)
doc.add_paragraph(", ".join(features))

doc.add_heading("Target Variable", level=1)
doc.add_paragraph("Ready_Status (1 = Ready, 0 = Not Ready)")

doc.add_heading("Readiness Logic", level=1)
doc.add_paragraph("If Level is JHS/SHS and FLT Overall Score is ≥ 75%, mark as Ready (1). Otherwise, Not Ready (0).")

doc.add_heading("Class Distribution", level=1)
for k, v in Counter(y).items():
    label = "Ready" if k == 1 else "Not Ready"
    doc.add_paragraph(f"{label}: {v} learners")

doc.add_heading("Model Used", level=1)
doc.add_paragraph("Support Vector Machine (SVM) with RBF kernel")

doc.add_heading("Model Performance", level=1)
doc.add_paragraph(f"Accuracy: {accuracy * 100:.2f}%")
doc.add_paragraph(f"Confusion Matrix:\n{conf_matrix}")
doc.add_paragraph(f"Classification Report:\n{class_report}")

doc.add_heading("Cross-Validation (5-Fold)", level=1)
doc.add_paragraph(f"Mean Accuracy: {cv_mean * 100:.2f}%")
doc.add_paragraph(f"Standard Deviation: {cv_std:.4f}")

doc.add_heading("Confusion Matrix", level=1)
doc.add_picture("C:\\Users\\User\\Desktop\\projects\\EALS - Copy\\backend\\models\\confusion_matrix.png")

doc.add_heading("ROC Curve", level=1)
doc.add_paragraph("The ROC curve shows the trade-off between sensitivity and specificity.")
doc.add_picture("C:\\Users\\User\\Desktop\\projects\\EALS - Copy\\backend\\models\\roc_curve.png")

doc.add_heading("Ready vs Not Ready Learners", level=1)
doc.add_picture("C:\\Users\\User\\Desktop\\projects\\EALS - Copy\\backend\\models\\ready_not_ready_graph.png")

doc.save("C:\\Users\\User\\Desktop\\projects\\EALS - Copy\\backend\\result\\model_analysis.docx")
print("📄 model_analysis.docx saved.")

# Save predictions and scores to Excel
test_results = X_test.copy()  # Input features for the test set
test_results['Actual-Result'] = y_test.values  # Add actual values
test_results['SVM-Predicted'] = y_pred  # Add SVM predicted values (1 or 0)
test_results['Probability-Ready'] = y_proba  # Add probabilities of being Ready (class 1)

# Column for SVM Result (Pass or Fail)
test_results['SVM-Result'] = test_results['SVM-Predicted'].apply(lambda x: 'Pass' if x == 1 else 'Fail')

# Save the dataframe to Excel
test_results.to_excel("C:\\Users\\User\\Desktop\\projects\\EALS - Copy\\backend\\result\\prediction_results.xlsx", index=False)

print("📊 prediction_results.xlsx saved.")

